from docx import Document
import re
from pathlib import Path

doc = Document(r"c:\Users\Yura\Downloads\Tymchenko.docx")
items = []
for p in doc.paragraphs:
    if p.text.strip():
        items.append(("p", p.text))
for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            s = cell.text.strip()
            if s:
                items.append((f"t{ti + 1}", s))

regex = re.compile(
    r"\b(я|мы|мой|моя|моё|мое|мои|меня|мне|мной|мною|"
    r"нам|нас|нами|наш|наша|наше|наши|нашего|нашей|нашему|нашим|наших|нашу|"
    r"моего|моей|моему|моим|моих|свой|своя|своё|свое|свои|своего|своей|своим)\b",
    re.IGNORECASE | re.UNICODE,
)

out = []
for loc, line in items:
    matches = list(regex.finditer(line))
    if matches:
        words = list(dict.fromkeys(m.group(0) for m in matches))
        out.append((loc, words, line))

path = Path(__file__).resolve().parent / "_tymchenko_first_person.txt"
with path.open("w", encoding="utf-8") as f:
    f.write(f"Найдено фрагментов: {len(out)}\n\n")
    for i, (loc, words, text) in enumerate(out, 1):
        f.write(f"[{i}] Слова: {', '.join(words)}\n")
        f.write(f"{text}\n\n")

print(path, len(out))
